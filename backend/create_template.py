from docx import Document

def create_default_template():
    doc = Document()
    
    # Add College and Department
    doc.add_heading('{{COLLEGE_NAME}}', 0)
    doc.add_heading('{{DEPARTMENT}}', 1)
    
    doc.add_paragraph('\n' * 5)
    
    # Add Report Topic
    doc.add_heading('REPORT ON: {{TOPIC}}', 1)
    
    doc.add_paragraph('\n' * 5)
    
    # Add Student Details
    doc.add_paragraph('Submitted by:')
    doc.add_paragraph('Name: {{STUDENT_NAME}}')
    doc.add_paragraph('Roll No: {{ROLL_NO}}')
    
    doc.add_page_break()
    
    # Sections
    sections = [
        ('Introduction', '{{INTRODUCTION}}'),
        ('Objectives', '{{OBJECTIVES}}'),
        ('Methodology', '{{METHODOLOGY}}'),
        ('Result', '{{RESULT}}'),
        ('Conclusion', '{{CONCLUSION}}'),
        ('References', '{{REFERENCES}}')
    ]
    
    for title, placeholder in sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph(placeholder)
        doc.add_paragraph('\n')

    doc.save('templates/default_template.docx')
    print("Default template created at backend/templates/default_template.docx")

if __name__ == "__main__":
    create_default_template()
