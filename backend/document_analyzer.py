import os
import uuid
import sys
sys.path.append(os.path.dirname(__file__))
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
from datetime import datetime
from models.analysis import SampleDocumentAnalysis

class DocumentAnalyzer:
    """Analyzes Word documents to extract formatting and structure information"""
    
    def __init__(self, uploads_dir: str = "uploads"):
        self.uploads_dir = uploads_dir
        os.makedirs(uploads_dir, exist_ok=True)
    
    def analyze_document(self, file_path: str, original_filename: str) -> SampleDocumentAnalysis:
        """
        Analyze a Word document and extract formatting information
        
        Args:
            file_path: Path to the uploaded document
            original_filename: Original filename
            
        Returns:
            SampleDocumentAnalysis object with detailed analysis
        """
        try:
            doc = Document(file_path)
            document_id = str(uuid.uuid4())
            
            # Extract basic information
            file_size = os.path.getsize(file_path)
            
            # Extract formatting information
            font_styles = self._extract_font_styles(doc)
            paragraph_styles = self._extract_paragraph_styles(doc)
            section_structure = self._analyze_section_structure(doc)
            header_footer_info = self._extract_header_footer(doc)
            page_setup = self._extract_page_setup(doc)
            
            # Identify placeholders and content structure
            identified_placeholders = self._identify_placeholders(doc)
            content_sections = self._identify_content_sections(doc)
            
            # Calculate formatting preservation score
            formatting_score = self._calculate_formatting_score(doc)
            
            # Template validation
            is_valid_template = self._validate_template(doc)
            template_compatibility = self._assess_template_compatibility(doc)
            recommended_placeholders = self._suggest_placeholders(doc)
            
            return SampleDocumentAnalysis(
                document_id=document_id,
                original_filename=original_filename,
                file_size=file_size,
                upload_timestamp=datetime.now(),
                font_styles=font_styles,
                paragraph_styles=paragraph_styles,
                section_structure=section_structure,
                header_footer_info=header_footer_info,
                page_setup=page_setup,
                identified_placeholders=identified_placeholders,
                content_sections=content_sections,
                formatting_preservation_score=formatting_score,
                is_valid_template=is_valid_template,
                template_compatibility=template_compatibility,
                recommended_placeholders=recommended_placeholders
            )
            
        except Exception as e:
            raise Exception(f"Document analysis failed: {str(e)}")
    
    def _extract_font_styles(self, doc: Document) -> Dict[str, Any]:
        """Extract font styling information from the document"""
        font_info = {
            'fonts_used': set(),
            'font_sizes': set(),
            'font_colors': set(),
            'bold_usage': 0,
            'italic_usage': 0,
            'underline_usage': 0
        }
        
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.font.name:
                    font_info['fonts_used'].add(run.font.name)
                if run.font.size:
                    font_info['font_sizes'].add(run.font.size.pt if hasattr(run.font.size, 'pt') else str(run.font.size))
                if run.font.color and run.font.color.rgb:
                    font_info['font_colors'].add(str(run.font.color.rgb))
                if run.bold:
                    font_info['bold_usage'] += 1
                if run.italic:
                    font_info['italic_usage'] += 1
                if run.underline:
                    font_info['underline_usage'] += 1
        
        # Convert sets to lists for JSON serialization
        font_info['fonts_used'] = list(font_info['fonts_used'])
        font_info['font_sizes'] = list(font_info['font_sizes'])
        font_info['font_colors'] = list(font_info['font_colors'])
        
        return font_info
    
    def _extract_paragraph_styles(self, doc: Document) -> Dict[str, Any]:
        """Extract paragraph styling information"""
        paragraph_info = {
            'alignment_types': set(),
            'spacing_before': [],
            'spacing_after': [],
            'line_spacing': set(),
            'styles_used': set()
        }
        
        for paragraph in doc.paragraphs:
            # Alignment
            alignment = paragraph.alignment
            if alignment:
                paragraph_info['alignment_types'].add(str(alignment))
            
            # Spacing
            if paragraph.paragraph_format.space_before:
                paragraph_info['spacing_before'].append(paragraph.paragraph_format.space_before.pt)
            if paragraph.paragraph_format.space_after:
                paragraph_info['spacing_after'].append(paragraph.paragraph_format.space_after.pt)
            
            # Line spacing
            if paragraph.paragraph_format.line_spacing:
                paragraph_info['line_spacing'].add(str(paragraph.paragraph_format.line_spacing))
            
            # Styles
            if paragraph.style:
                paragraph_info['styles_used'].add(paragraph.style.name)
        
        # Convert sets to lists
        paragraph_info['alignment_types'] = list(paragraph_info['alignment_types'])
        paragraph_info['line_spacing'] = list(paragraph_info['line_spacing'])
        paragraph_info['styles_used'] = list(paragraph_info['styles_used'])
        
        return paragraph_info
    
    def _analyze_section_structure(self, doc: Document) -> List[Dict[str, Any]]:
        """Analyze the document's section structure"""
        sections = []
        
        for i, section in enumerate(doc.sections):
            section_info = {
                'section_number': i + 1,
                'start_page': 'auto',  # Would need more complex analysis
                'page_orientation': 'portrait' if section.orientation == 0 else 'landscape',
                'page_width': section.page_width.pt if hasattr(section.page_width, 'pt') else str(section.page_width),
                'page_height': section.page_height.pt if hasattr(section.page_height, 'pt') else str(section.page_height),
                'margins': {
                    'top': section.top_margin.pt if hasattr(section.top_margin, 'pt') else str(section.top_margin),
                    'bottom': section.bottom_margin.pt if hasattr(section.bottom_margin, 'pt') else str(section.bottom_margin),
                    'left': section.left_margin.pt if hasattr(section.left_margin, 'pt') else str(section.left_margin),
                    'right': section.right_margin.pt if hasattr(section.right_margin, 'pt') else str(section.right_margin)
                }
            }
            sections.append(section_info)
        
        return sections
    
    def _extract_header_footer(self, doc: Document) -> Dict[str, Any]:
        """Extract header and footer information"""
        header_footer_info = {
            'has_header': False,
            'has_footer': False,
            'header_content': [],
            'footer_content': [],
            'header_font_info': {},
            'footer_font_info': {}
        }
        
        # Check first section for headers/footers
        if doc.sections:
            section = doc.sections[0]
            
            # Check headers
            if section.header.paragraphs:
                header_footer_info['has_header'] = True
                for para in section.header.paragraphs:
                    if para.text.strip():
                        header_footer_info['header_content'].append(para.text)
            
            # Check footers
            if section.footer.paragraphs:
                header_footer_info['has_footer'] = True
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        header_footer_info['footer_content'].append(para.text)
        
        return header_footer_info
    
    def _extract_page_setup(self, doc: Document) -> Dict[str, Any]:
        """Extract page setup information"""
        if not doc.sections:
            return {}
        
        section = doc.sections[0]
        return {
            'page_width': section.page_width.pt if hasattr(section.page_width, 'pt') else str(section.page_width),
            'page_height': section.page_height.pt if hasattr(section.page_height, 'pt') else str(section.page_height),
            'orientation': 'portrait' if section.orientation == 0 else 'landscape',
            'margins': {
                'top': section.top_margin.pt if hasattr(section.top_margin, 'pt') else 'default',
                'bottom': section.bottom_margin.pt if hasattr(section.bottom_margin, 'pt') else 'default',
                'left': section.left_margin.pt if hasattr(section.left_margin, 'pt') else 'default',
                'right': section.right_margin.pt if hasattr(section.right_margin, 'pt') else 'default'
            }
        }
    
    def _identify_placeholders(self, doc: Document) -> List[str]:
        """Identify placeholder patterns in the document"""
        placeholders = set()
        placeholder_pattern = r'\{\{[^}]+\}\}'
        
        for paragraph in doc.paragraphs:
            matches = re.findall(placeholder_pattern, paragraph.text)
            placeholders.update([match.strip('{}') for match in matches])
        
        # Also check in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = re.findall(placeholder_pattern, paragraph.text)
                        placeholders.update([match.strip('{}') for match in matches])
        
        return list(placeholders)
    
    def _identify_content_sections(self, doc: Document) -> List[str]:
        """Identify content sections based on headings"""
        sections = []
        heading_patterns = [
            r'^\s*introduction\s*$', r'^\s*objective[s]?\s*$', 
            r'^\s*methodology\s*$', r'^\s*result[s]?\s*$',
            r'^\s*conclusion\s*$', r'^\s*reference[s]?\s*$'
        ]
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.lower().strip()
            for pattern in heading_patterns:
                if re.match(pattern, text, re.IGNORECASE):
                    sections.append(text.title())
                    break
        
        return list(set(sections))  # Remove duplicates
    
    def _calculate_formatting_score(self, doc: Document) -> float:
        """Calculate a formatting preservation score (0-100)"""
        score = 100.0
        
        # Deduct points for inconsistent formatting
        font_styles = self._extract_font_styles(doc)
        paragraph_styles = self._extract_paragraph_styles(doc)
        
        # Check for excessive font variations
        if len(font_styles['fonts_used']) > 5:
            score -= 10
        
        # Check for consistent paragraph styles
        if len(paragraph_styles['styles_used']) > 10:
            score -= 15
        
        # Check for proper structure
        sections = self._identify_content_sections(doc)
        if len(sections) < 3:
            score -= 20
        
        return max(0, score)
    
    def _validate_template(self, doc: Document) -> bool:
        """Validate if document can be used as a template"""
        # Basic validation criteria
        has_content = len(doc.paragraphs) > 5
        has_structure = len(self._identify_content_sections(doc)) >= 2
        has_placeholders = len(self._identify_placeholders(doc)) > 0
        
        return has_content and has_structure and has_placeholders
    
    def _assess_template_compatibility(self, doc: Document) -> str:
        """Assess template compatibility level"""
        score = self._calculate_formatting_score(doc)
        
        if score >= 80:
            return "high"
        elif score >= 60:
            return "medium"
        elif score >= 40:
            return "low"
        else:
            return "incompatible"
    
    def _suggest_placeholders(self, doc: Document) -> List[str]:
        """Suggest standard placeholders based on document content"""
        suggested = [
            'STUDENT_NAME', 'ROLL_NO', 'TOPIC', 'COLLEGE_NAME', 
            'DEPARTMENT', 'INTRODUCTION', 'OBJECTIVES', 
            'METHODOLOGY', 'RESULT', 'CONCLUSION', 'REFERENCES'
        ]
        
        # Add any existing placeholders found
        existing = self._identify_placeholders(doc)
        suggested.extend([p for p in existing if p not in suggested])
        
        return suggested[:15]  # Limit to 15 suggestions